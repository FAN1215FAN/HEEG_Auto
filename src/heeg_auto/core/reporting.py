from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

from heeg_auto.config.settings import REPORT_DIR, ensure_artifact_dirs

REPORT_TITLE = "HEEG 自动化测试执行报告"


def build_report_base_name(report_timestamp: str) -> str:
    return f"HEEG_Auto_Report_{report_timestamp}"


def generate_reports(result: dict[str, Any]) -> dict[str, str]:
    ensure_artifact_dirs()
    report_timestamp = str(result.get("report_timestamp") or "unknown")
    base_name = build_report_base_name(report_timestamp)
    json_path = REPORT_DIR / f"{base_name}.json"
    docx_path = REPORT_DIR / f"{base_name}.docx"

    payload = _build_report_payload(result=result, json_path=json_path, docx_path=docx_path)
    json_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    _write_docx_report(payload, docx_path)

    return {"json_path": str(json_path), "docx_path": str(docx_path)}


def _build_report_payload(result: dict[str, Any], json_path: Path, docx_path: Path) -> dict[str, Any]:
    screenshot_paths = [str(path) for path in result.get("artifact_paths", [])]
    return {
        "report_title": REPORT_TITLE,
        "report_timestamp": result.get("report_timestamp", ""),
        "report_files": {
            "json": str(json_path),
            "docx": str(docx_path),
        },
        "case_id": result.get("case_id", ""),
        "case_name": result.get("case_name", ""),
        "case_path": result.get("case_path", ""),
        "tags": result.get("tags", []),
        "module_chain_labels": result.get("module_chain_labels", []),
        "status": result.get("status", "FAIL"),
        "passed": bool(result.get("passed", False)),
        "started_at": result.get("started_at", ""),
        "finished_at": result.get("finished_at", ""),
        "duration_seconds": result.get("duration_seconds", 0),
        "context": result.get("context", {}),
        "module_results": result.get("module_results", []),
        "error_summary": result.get("error_summary", ""),
        "error_detail": result.get("error_detail", ""),
        "failure": result.get("failure", {}),
        "artifact_paths": screenshot_paths,
        "body_screenshot_paths": _pick_body_screenshots(screenshot_paths),
        "environment": result.get("environment", {}),
    }


def _pick_body_screenshots(screenshot_paths: list[str]) -> list[str]:
    if not screenshot_paths:
        return []

    preferred: list[str] = []
    for keyword in ("_screen.png", "_active.png"):
        for path in screenshot_paths:
            if path.endswith(keyword) and path not in preferred:
                preferred.append(path)
                break

    if len(preferred) < 2:
        for path in screenshot_paths:
            if path not in preferred:
                preferred.append(path)
            if len(preferred) >= 2:
                break

    return preferred[:2]


def _write_docx_report(payload: dict[str, Any], docx_path: Path) -> None:
    document = Document()
    _set_default_style(document)

    title = document.add_heading(payload["report_title"], level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    summary_lines = [
        f"执行结果：{'通过' if payload['passed'] else '失败'}",
        f"用例编号：{payload.get('case_id', '-')}",
        f"用例名称：{payload['case_name']}",
        f"模块链：{' -> '.join(payload.get('module_chain_labels', [])) or '-'}",
        f"执行时间：{payload['started_at']} 至 {payload['finished_at']}",
        f"总耗时：{payload['duration_seconds']} 秒",
    ]
    _add_bullet_lines(document, "执行摘要", summary_lines)

    basic_info = [
        ("应用路径", payload["environment"].get("app_path", "")),
        ("用例文件", payload.get("case_path", "")),
        ("工作目录", payload["environment"].get("cwd", "")),
        ("Python 版本", payload["environment"].get("python_version", "")),
        ("开始时间", payload.get("started_at", "")),
        ("结束时间", payload.get("finished_at", "")),
        ("执行耗时（秒）", str(payload.get("duration_seconds", ""))),
    ]
    _add_key_value_table(document, "基本信息", basic_info)

    case_info = [
        ("用例编号", payload.get("case_id", "")),
        ("用例名称", payload.get("case_name", "")),
        ("标签", "、".join(payload.get("tags", [])) or "-"),
        ("模块链", " -> ".join(payload.get("module_chain_labels", [])) or "-"),
        ("执行状态", "PASS" if payload.get("passed") else "FAIL"),
        ("Word 报告", payload["report_files"].get("docx", "")),
    ]
    _add_key_value_table(document, "用例执行结果", case_info)

    _add_module_summary_table(document, payload.get("module_results", []))
    _add_failure_section(document, payload)
    _add_screenshot_section(document, payload)
    _add_appendix(document, payload)

    document.save(docx_path)


def _set_default_style(document: Document) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Microsoft YaHei"
    normal_style._element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
        "Microsoft YaHei",
    )


def _add_bullet_lines(document: Document, heading: str, lines: list[str]) -> None:
    document.add_heading(heading, level=1)
    for line in lines:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(line)


def _add_key_value_table(document: Document, heading: str, rows: list[tuple[str, str]]) -> None:
    document.add_heading(heading, level=1)
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "字段"
    header_cells[1].text = "内容"
    for key, value in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = value or "-"


def _add_module_summary_table(document: Document, module_results: list[dict[str, Any]]) -> None:
    document.add_heading("模块执行摘要", level=1)
    table = document.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "序号"
    header[1].text = "模块标识"
    header[2].text = "模块名称"
    header[3].text = "预期状态"
    header[4].text = "执行结果"
    header[5].text = "失败步骤"

    for index, module in enumerate(module_results, start=1):
        row = table.add_row().cells
        row[0].text = str(index)
        row[1].text = module.get("module_id", "") or "-"
        row[2].text = module.get("module_label", "") or "-"
        row[3].text = module.get("expected_status", "") or "-"
        row[4].text = f"{module.get('status', '-')}\n耗时：{module.get('duration_seconds', 0)} 秒"
        row[5].text = module.get("failed_step", "") or "-"


def _add_failure_section(document: Document, payload: dict[str, Any]) -> None:
    document.add_heading("故障详情", level=1)
    if payload.get("passed"):
        document.add_paragraph("本次执行全部通过，无故障明细。")
        return

    paragraphs = [
        f"错误摘要：{payload.get('error_summary') or '-'}",
        f"失败模块：{payload.get('failure', {}).get('module_id') or '-'}",
        f"失败步骤：{payload.get('failure', {}).get('failed_step') or '-'}",
        f"故障截图数量：{len(payload.get('artifact_paths', []))}",
    ]
    for line in paragraphs:
        document.add_paragraph(line)


def _add_screenshot_section(document: Document, payload: dict[str, Any]) -> None:
    document.add_heading("故障截图", level=1)
    screenshot_paths = payload.get("body_screenshot_paths", [])
    if not screenshot_paths:
        document.add_paragraph("当前没有可插入正文的截图。")
        return

    for image_path in screenshot_paths:
        path = Path(image_path)
        if not path.exists():
            continue
        document.add_paragraph(path.name)
        document.add_picture(str(path), width=Inches(6.2))
        document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def _add_appendix(document: Document, payload: dict[str, Any]) -> None:
    document.add_heading("附录：原始结果与错误堆栈", level=1)
    if payload.get("module_results"):
        document.add_paragraph("模块原始执行结果：")
        document.add_paragraph(json.dumps(payload["module_results"], ensure_ascii=False, indent=2))
    document.add_paragraph("全部截图路径：")
    for path in payload.get("artifact_paths", []):
        document.add_paragraph(path, style="List Bullet")

    document.add_paragraph(f"JSON 报告：{payload['report_files'].get('json', '')}")
    document.add_paragraph(f"Word 报告：{payload['report_files'].get('docx', '')}")

    stack_trace = payload.get("error_detail") or "-"
    document.add_paragraph("完整错误堆栈：")
    document.add_paragraph(stack_trace)